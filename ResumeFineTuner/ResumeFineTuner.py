from docx import Document
from datetime import date
import re
import os
import openai
import json

class ResumeFineTuner():

    def __init__(self, job_descr=''):
        self.job_description = job_descr
        with open("data/full_experience.txt", "r") as file:
            self.actual_experience = file.read()
        with open("data/example_profile.txt", "r") as file:
            self.example = file.read()
        with open("data/key.txt", "r") as file:
            self.key = file.read()

    def fine_tune_cv(self):
        self.generate_prompts()
        self.generate_sections()
        self.parse_output()
        self.update_document()

    def generate_prompts(self):
        print("---------------------------------------------------")
        print("PREPARING PROMPTS...")
        print("---------------------------------------------------")

        self.system_prompt = '''You are a resume expert and you will write the section of a resume based on actual experience and tuned for the job description that the user asks.
You will produce the following outputs:
#PROFILE: high level presentation of candidate tuned for the job description (max 100 words).
#MANAGER PROJECTS: select and rephrase main projects inherent for the job description among the manager ones (max 3 points).
#SENIOR DATA SCIENTIST PROJECTS: select and rephrase main projects inherent for the job description among the senior data scientist ones (max 3 points).
#RESEARCH FELLOW PROJECTS: select and rephrase main projects inherent for the job description among the research fellow ones (max 3 points).
#KEY WORDS: key words related to actual experience and tuned for the job description (max 5 keywords).

The candidate's actual experience is the following:
{}

Note: Do not invet informations, just rephrase actual experience to meet the job description. Create the output in English.
'''.format(self.actual_experience)
        print(self.system_prompt)
        self.assistant_prompt = '''Here an example of the desired output:
{}
'''.format(self.example)
        print(self.assistant_prompt)
        self.user_prompt = '''
Write the candidate's resume in English for the following job description:
{}
'''.format(self.job_description)
        print(self.user_prompt)

    def generate_sections(self):
        print("---------------------------------------------------")
        print("GENERATING CV SECTIONS...")
        print("---------------------------------------------------")

        openai.api_key = self.key
        messages = [
            {"role": "system", "content": self.system_prompt},
            {"role": "assistant", "content": self.assistant_prompt},
            {"role": "user", "content": self.user_prompt}
        ]
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=messages,
            temperature=0.5,
            max_tokens=1024
        )
        self.output = response['choices'][0]['message']['content']
        print(self.output)

        print("---------------------------------------------------")
        print("OUTPUT GENERATED!")
        print("---------------------------------------------------")

    def parse_output(self):
        print("---------------------------------------------------")
        print("PARSING OUTPUT...")
        print("---------------------------------------------------")

        splits = ['#PROFILE:', '#MANAGER PROJECTS:', '#SENIOR DATA SCIENTIST PROJECTS:', 
            '#RESEARCH FELLOW PROJECTS:', '#KEY WORDS:']
        for split in splits:
            self.output = self.output.replace(split, '###')
        output_list = self.output.split('###')[1:]

        out_dict = {}
        for split, o in zip(splits, output_list):
            section_name = split.replace(':', '').replace(' ', '').replace('\n', '')
            if '*' in o:
                subsplits = o.split('*')[1:]
                for i, subsplit in enumerate(subsplits):
                    out_dict[section_name+str(i+1)] = subsplit.replace('\n', '').lstrip().rstrip()
            else:
                out_dict[section_name] = o.replace('\n', '').lstrip()
        self.out_dict = out_dict

        print("---------------------------------------------------")
        print("PARSED OUTPUT!")
        print("---------------------------------------------------")
        print(json.dumps(out_dict, indent=1))
        

    def update_document(self):

        def docx_replace_regex(doc_obj, regex, replace):

            for p in doc_obj.paragraphs:
                if regex in p.text:
                    p.text = p.text.replace(regex, replace)

            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        docx_replace_regex(cell, regex, replace)

        print("---------------------------------------------------")
        print("UPDATING CV...")
        print("---------------------------------------------------")

        document = Document('data/CV_template.docx')
        for key in self.out_dict:
            replace_text = self.out_dict[key]
            docx_replace_regex(document, key, replace_text)

        document.save('outputs/CV_APazzaglia_{}.docx'.format(date.today()))

        print("---------------------------------------------------")
        print("CV UPDATED!")
        print("---------------------------------------------------")